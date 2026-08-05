"""El instrumento de la no contaminación (REQ-48).

Compara la huella de texto de los documentos de control —39056 (Big Data) y 39062
(Patrones)— contra el registro versionado en `pruebas/huellas.yaml`. Es el instrumento
que heredan las cinco fases siguientes del milestone v2.0: regenerar los cuatro
documentos de control después de cada fase debe dejar su texto, su informe de
validación y la forma de su `MANIFIESTO.yaml` idénticos a los registrados.

No cuelga del ciclo de pruebas unitarias (D-18): generar dos cursos completos con Word
sería lento y las pruebas de `pruebas/` deben seguir corriendo en segundos. Se corre a
mano al cerrar cada fase:

    python src/huella.py verificar    # compara contra el registro
    python src/huella.py registrar    # acepta el estado actual como línea base
"""

from __future__ import annotations

import hashlib
import sys
from datetime import date
from pathlib import Path

import yaml
from docx import Document
from docx.oxml.ns import qn

import generar

RAIZ = Path(__file__).resolve().parent.parent
REGISTRO = RAIZ / "pruebas" / "huellas.yaml"

# Los documentos de control son los de Adrian (D-22/D-25). 38985 no entra: es el curso
# que va a cambiar en la Fase 14.
CONTROL = (
    ("cursos/2026-2/39056-big-data/curso.yaml", ("961", "962")),
    ("cursos/2026-2/39062-patrones-de-comportamiento/curso.yaml", ("971", "972")),
)

CABECERA = (
    "# Huellas de texto de los documentos de control (REQ-48). Lo genera\n"
    "# src/huella.py: no lo edites a mano.\n"
    "#\n"
    "#     python src/huella.py verificar    # compara contra este registro\n"
    "#     python src/huella.py registrar    # acepta el estado actual como línea base\n"
    "#\n"
    "# Tres hashes por documento:\n"
    "#   texto_docx  el texto extraído del .docx — no su binario, cuyo zip lleva marcas\n"
    "#               de tiempo y cambiaría en cada corrida.\n"
    "#   informe     la salida de validar.py. Es por curso, así que se repite entre los\n"
    "#               grupos de un mismo curso: es correcto, no está duplicado por error.\n"
    "#   manifiesto  la FORMA del MANIFIESTO.yaml: qué claves emite y con qué estructura,\n"
    "#               sin `generado`, `commit` ni el sha256/bytes de los archivos, que\n"
    "#               cambian en cada corrida por diseño. También es por curso.\n"
)


class ErrorHuella(Exception):
    """Falta el registro, o un documento de control no se pudo generar."""


def _texto_de(elemento) -> str:
    """Todo el texto (`<w:t>`) dentro de un elemento OOXML, en orden de aparición."""
    return "".join(t.text or "" for t in elemento.iter(qn("w:t")))


def extraer_texto(ruta_docx: Path | str) -> str:
    """El texto del documento en orden real: encabezado, cuerpo (párrafos y tablas), pie.

    Recorre el XML crudo en vez de `doc.paragraphs`/`doc.tables`, que son listas
    separadas y no conservan el orden, y en vez de `row.cells[i].text`, que resuelve
    las celdas con `vMerge` duplicando el texto de la celda de origen.
    """
    doc = Document(str(ruta_docx))
    bloques: list[str] = []
    for sec in doc.sections:          # hoy están vacíos; no cuesta nada dejarlo listo
        for p in sec.header.paragraphs:
            if p.text.strip():
                bloques.append(p.text)
    for hijo in doc.element.body.iterchildren():
        tag = hijo.tag.rsplit("}", 1)[-1]
        if tag == "p":
            bloques.append(_texto_de(hijo))
        elif tag == "tbl":
            for tr in hijo.iter(qn("w:tr")):
                bloques.append(" | ".join(_texto_de(tc) for tc in tr.findall(qn("w:tc"))))
    for sec in doc.sections:
        for p in sec.footer.paragraphs:
            if p.text.strip():
                bloques.append(p.text)
    return "\n".join(bloques)


def sha_texto(texto: str) -> str:
    return hashlib.sha256(texto.encode("utf-8")).hexdigest()


# Lo que cambia en cada corrida por diseño y no dice nada sobre la forma (D-27).
VOLATILES_MANIFIESTO = ("generado", "commit")
VOLATILES_ARCHIVO = ("sha256", "bytes")


def forma_del_manifiesto(ruta: Path | str) -> str:
    """sha256 del MANIFIESTO.yaml con sus claves volátiles fuera.

    Vigila qué claves emite y con qué estructura, no sus valores por corrida. Si una fase
    posterior añadiera un bloque al manifiesto —el segundo nivel de la calificación, la
    rúbrica— este hash lo denuncia, que es la mitad de REQ-48 que la huella de texto no ve.
    """
    datos = yaml.safe_load(Path(ruta).read_text(encoding="utf-8")) or {}
    for clave in VOLATILES_MANIFIESTO:
        datos.pop(clave, None)
    for archivo in datos.get("archivos", []):
        for clave in VOLATILES_ARCHIVO:
            archivo.pop(clave, None)
    return sha_texto(
        yaml.safe_dump(datos, allow_unicode=True, sort_keys=False, default_flow_style=False)
    )


def cargar() -> dict[str, dict]:
    """Lee `pruebas/huellas.yaml`. Devuelve {} si aún no existe."""
    if not REGISTRO.exists():
        return {}
    datos = yaml.safe_load(REGISTRO.read_text(encoding="utf-8")) or {}
    return datos.get("documentos", {})


def guardar(documentos: dict[str, dict]) -> None:
    cuerpo = {"documentos": {k: documentos[k] for k in sorted(documentos)}}
    texto = yaml.safe_dump(cuerpo, allow_unicode=True, sort_keys=False, default_flow_style=False)
    REGISTRO.write_text(CABECERA + texto, encoding="utf-8")


def _generar_control(restaurar_manifiesto: bool) -> dict[str, dict]:
    """Regenera los .docx de control y devuelve sus huellas, indexadas por ciclo:clave:grupo."""
    actuales: dict[str, dict] = {}
    for rel, grupos in CONTROL:
        ruta = RAIZ / rel
        if not ruta.exists():
            raise ErrorHuella(f"Falta el curso de control {rel}.")
        manifiesto = ruta.parent / "MANIFIESTO.yaml"
        previo = manifiesto.read_bytes() if manifiesto.exists() else None
        try:
            paq = generar.paquete(ruta, pdf=False, grupos=list(grupos))
            # El manifiesto RECIÉN generado, antes de que el finally restaure el anterior.
            forma = forma_del_manifiesto(paq.manifiesto)
        except generar.ErrorGenerar as e:
            raise ErrorHuella(f"{rel}: no se pudo generar — {e}") from e
        finally:
            if restaurar_manifiesto:
                if previo is not None:
                    manifiesto.write_bytes(previo)
                else:
                    manifiesto.unlink(missing_ok=True)
        informe = sha_texto(paq.informe.texto())
        for a in paq.archivos:
            if a.suffix != ".docx":
                continue
            grupo = a.stem.rsplit("-", 1)[-1]
            actuales[f"{paq.curso.ciclo}:{paq.curso.clave}:{grupo}"] = {
                "ciclo": paq.curso.ciclo,
                "clave": paq.curso.clave,
                "grupo": grupo,
                "texto_docx": sha_texto(extraer_texto(a)),
                "informe": informe,
                "manifiesto": forma,
            }
    return actuales


def _rotulo(clave: str) -> str:
    _, clave_materia, grupo = clave.split(":")
    return f"{clave_materia} grupo {grupo}"


CAMPOS = ("texto_docx", "informe", "manifiesto")
QUE_CAMBIO = {
    "texto_docx": "el texto del documento",
    "informe": "el informe de validación",
    "manifiesto": "la forma del MANIFIESTO.yaml",
}


def verificar() -> tuple[list[str], list[str]]:
    """Compara el estado actual contra el registro. No escribe nada (D-28)."""
    registradas = cargar()
    if not registradas:
        raise ErrorHuella(
            "No hay huellas registradas todavía. Registra la línea base con "
            "`python src/huella.py registrar` antes de comparar."
        )
    actuales = _generar_control(restaurar_manifiesto=True)
    problemas: list[str] = []
    intactos: list[str] = []
    for clave in sorted(set(registradas) | set(actuales)):
        vieja, nueva = registradas.get(clave), actuales.get(clave)
        if nueva is None:
            problemas.append(
                f"{_rotulo(clave)}: está registrado pero ya no se generó. "
                f"¿Cambió la lista de grupos del curso? Acepta el cambio con "
                f"`python src/huella.py registrar`."
            )
            continue
        if vieja is None:
            problemas.append(
                f"{_rotulo(clave)}: documento nuevo, sin huella registrada. "
                f"Regístralo con `python src/huella.py registrar`."
            )
            continue
        cambios = [c for c in CAMPOS if vieja.get(c) != nueva[c]]
        if not cambios:
            intactos.append(clave)
            continue
        for c in cambios:
            que = QUE_CAMBIO[c]
            problemas.append(
                f"{_rotulo(clave)}: cambió {que} ({vieja.get(c, '—')[:12]} → {nueva[c][:12]}). "
                f"Mira qué lo movió con `git diff` y, si el cambio es deliberado, acéptalo con "
                f"`python src/huella.py registrar`."
            )
    return problemas, intactos


def registrar() -> list[str]:
    """Acepta el estado actual como línea base.

    A diferencia de `verificar`, **sí** deja escritos los MANIFIESTO.yaml de los cursos de
    control: aquí el cambio es deliberado y debe quedar en el commit (D-23/D-28).
    """
    previas = cargar()
    actuales = _generar_control(restaurar_manifiesto=False)
    hoy = date.today().isoformat()
    lineas: list[str] = []
    for clave in sorted(actuales):
        nueva, vieja = actuales[clave], previas.get(clave)
        igual = vieja is not None and all(vieja.get(c) == nueva[c] for c in CAMPOS)
        nueva["registrado"] = vieja["registrado"] if igual and vieja.get("registrado") else hoy
        if vieja is None:
            lineas.append(f"+ {_rotulo(clave)}   huella registrada ({nueva['texto_docx'][:12]})")
        elif igual:
            lineas.append(f"= {_rotulo(clave)}   sin cambios")
        else:
            lineas.append(
                f"! {_rotulo(clave)}   {vieja['texto_docx'][:12]} → {nueva['texto_docx'][:12]}"
            )
    for clave in sorted(set(previas) - set(actuales)):
        lineas.append(f"- {_rotulo(clave)}   ya no se genera; se retira del registro")
    guardar(actuales)
    return lineas


def main(argv: list[str]) -> int:
    orden = argv[1] if len(argv) > 1 else "verificar"
    try:
        if orden == "registrar":
            for linea in registrar():
                print(f"  {linea}")
            print(f"\nRegistro escrito en {REGISTRO.relative_to(RAIZ).as_posix()}.")
            print("Los MANIFIESTO.yaml de los cursos de control quedaron reescritos: van en el commit.")
        elif orden == "verificar":
            problemas, intactos = verificar()
            for clave in intactos:
                print(f"  ✓ {_rotulo(clave):<20} huella intacta")
            if problemas:
                for p in problemas:
                    print(f"  ✗ {p}", file=sys.stderr)
                print(f"\n{len(problemas)} problema(s).", file=sys.stderr)
                return 1
            print(f"\nTodo intacto. {len(intactos)} documentos comparados.")
        else:
            print(f"Orden desconocida: {orden}", file=sys.stderr)
            print("Uso: python src/huella.py [verificar|registrar]", file=sys.stderr)
            return 2
    except ErrorHuella as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
