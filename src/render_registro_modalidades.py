"""Rellena el Registro de Modalidades de Acreditación Diversas desde YAML.

El formato se mantiene bajo la misma custodia que las plantillas CIAD: primero se
obtiene una copia verificada y fresca, luego solo se escriben sus celdas de datos.
La plantilla registrada nunca se modifica.

Uso:
    python src/render_registro_modalidades.py <registro.yaml> <salida.docx>
"""

from __future__ import annotations

import sys
import tempfile
from datetime import date, timedelta
from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

from calendario import ErrorCalendario, Periodo, cargar
from plantillas import copia_de_trabajo

PLANTILLA = "registro_modalidades_acreditacion_diversas"

MESES = (
    "enero",
    "febrero",
    "marzo",
    "abril",
    "mayo",
    "junio",
    "julio",
    "agosto",
    "septiembre",
    "octubre",
    "noviembre",
    "diciembre",
)


class ErrorRegistro(ValueError):
    """Los datos no pueden representarse fielmente en el formato institucional."""


def _texto(valor: object) -> str:
    return "" if valor is None else str(valor)


def _reemplazar(parrafo, valor: object) -> None:
    """Cambia solo los runs y conserva las propiedades del párrafo de la plantilla."""
    for nodo in list(parrafo._p):
        if nodo.tag == qn("w:r"):
            parrafo._p.remove(nodo)
    parrafo.add_run(_texto(valor))


def _celda(celda, valor: object) -> None:
    _reemplazar(celda.paragraphs[0], valor)


def _campo(tabla, fila: int, columna: int, valor: object) -> None:
    _celda(tabla.cell(fila, columna), valor)


def _recortar_filas(tabla, desde: int) -> None:
    """Elimina las filas de captura que el registro no utiliza.

    Mantener las filas vacías de la plantilla puede desplazar el bloque de firmas a
    la página siguiente y superponerlo con la Sección II. Los encabezados y las filas
    que contienen datos se conservan intactos.
    """
    for fila in list(tabla.rows[desde:]):
        tabla._tbl.remove(fila._tr)


def _texto_xml(parrafo) -> str:
    return "".join(nodo.text or "" for nodo in parrafo.iter(qn("w:t")))


def _reemplazar_xml(parrafo_xml, valor: str) -> None:
    """Escribe en un cuadro de texto sin alterar su geometría ni sus párrafos."""
    for nodo in list(parrafo_xml):
        if nodo.tag == qn("w:r"):
            parrafo_xml.remove(nodo)
    parrafo = Paragraph(parrafo_xml, None)
    parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = parrafo.add_run(valor)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.italic = False
    run.font.bold = False


def _salto_antes_de_actividades(documento: Document) -> None:
    """Mantiene las actividades en la tercera página que dispone el formato."""
    for parrafo in documento.element.body.iter(qn("w:p")):
        texto = _texto_xml(parrafo)
        if texto.count("IV. Actividades a realizar") == 2:
            # El ancla del cuadro es flotante y Word ignora el salto dentro del
            # mismo párrafo. Un párrafo independiente sí fija el ancla en página 3.
            separador = OxmlElement("w:p")
            run = OxmlElement("w:r")
            salto = OxmlElement("w:br")
            salto.set(qn("w:type"), "page")
            run.append(salto)
            separador.append(run)
            parrafo.addprevious(separador)
            return
    raise ErrorRegistro("No se localizó el salto de página previo a las actividades.")


def _fecha_larga(valor: date) -> str:
    return f"{valor.day} de {MESES[valor.month - 1]} de {valor.year}"


def _rango_actividad(actividad: dict, ciclo: str) -> str:
    calendario = cargar(ciclo)

    if periodo := actividad.get("periodo_calendario"):
        datos = calendario.periodos.get(periodo)
        if not isinstance(datos, Periodo):
            raise ErrorRegistro(f"Periodo inexistente en el calendario {ciclo}: {periodo}")
        inicio, fin = datos.inicio, datos.fin
    else:
        inicio = calendario.semana(actividad["inicio_semana"]).inicio
        fin = calendario.semana(actividad["fin_semana"]).fin
        while calendario.es_suspension(inicio):
            inicio += timedelta(days=1)

    return f"Del {_fecha_larga(inicio)} al {_fecha_larga(fin)}"


def _actividades_texto(actividades: list[dict], ciclo: str) -> str:
    lineas = []
    for actividad in actividades:
        rango = _rango_actividad(actividad, ciclo)
        producto = f" Entrega: {actividad['producto']}." if actividad.get("producto") else ""
        lineas.append(f"{rango}. {actividad['descripcion']}.{producto}")
    return "\n".join(lineas)


def _rellenar_cuerpo(documento: Document, registro: dict) -> None:
    marcadores = {
        "Mencionar el motivo por el cuál se ha elegido esta modalidad.": registro["justificacion"],
        (
            "Puede ser solo una competencia que cumpla con responder a ¿Qué hará el alumno? ¿Cómo "
            "lo hará? ¿Para qué? Y ¿Con qué actitud y/o valor?"
        ): registro["competencias"],
        (
            "Preferentemente, incluir la competencia general del Programa de Unidad de Aprendizaje y, "
            "de manera opcional incluir las competencias de cada Unidad Temática."
        ): "",
        (
            "Se recomienda enlistar cada actividad a realizar con las fechas correspondientes a la entrega "
            "de cada producto solicitado. Si lo considera pertinente, podrá organizar las actividades de "
            "acuerdo a cada Unidad temática del Programa de Unidad de Aprendizaje."
        ): _actividades_texto(registro["actividades"], registro["ciclo"]),
        (
            "Son las referencias principales a utilizar por el alumno, preferentemente 3 y en formato APA. "
            "Utilizar las referencias que se incluyen en el Programa de Unidad de Aprendizaje."
        ): "\n".join(registro["referencias"]),
    }
    reemplazos = {texto: 0 for texto in marcadores}
    for parrafo in documento.element.body.iter(qn("w:p")):
        texto = _texto_xml(parrafo)
        if texto in marcadores:
            _reemplazar_xml(parrafo, marcadores[texto])
            reemplazos[texto] += 1
        elif texto.strip() == "V. Evaluación y acreditación":
            padre = parrafo.getparent()
            siguiente = padre[padre.index(parrafo) + 1]
            _reemplazar_xml(siguiente, registro["evaluacion_acreditacion"])

    faltantes = [texto for texto, total in reemplazos.items() if total == 0]
    if faltantes:
        raise ErrorRegistro("No se localizaron todos los cuadros de texto del formato.")
    _salto_antes_de_actividades(documento)


def _validar(datos: dict) -> None:
    if not isinstance(datos, dict):
        raise ErrorRegistro("El archivo debe contener un objeto registro.")
    registro = datos.get("registro") or {}
    if not isinstance(registro, dict):
        raise ErrorRegistro("registro debe declararse como campos.")
    requeridos = (
        "unidad_academica",
        "periodo_estudio",
        "tipo_modalidad",
        "clave",
        "nombre_modalidad",
        "plan_estudios",
        "etapa_formacion",
        "programas",
        "creditos",
        "estudiantes",
        "responsables",
        "ciclo",
        "justificacion",
        "competencias",
        "actividades",
        "evaluacion_acreditacion",
        "referencias",
    )
    faltan = [campo for campo in requeridos if not registro.get(campo)]
    if faltan:
        raise ErrorRegistro("Faltan datos obligatorios: " + ", ".join(faltan))

    if str(registro["periodo_estudio"]) != str(registro["ciclo"]):
        raise ErrorRegistro("El periodo de estudio debe coincidir con el ciclo del calendario.")

    try:
        calendario = cargar(str(registro["ciclo"]))
    except ErrorCalendario as error:
        raise ErrorRegistro(str(error)) from error

    for campo in ("programas", "referencias", "actividades"):
        if not isinstance(registro[campo], list):
            raise ErrorRegistro(f"{campo} debe declararse como lista.")
    for campo in ("programas", "referencias"):
        if any(not isinstance(valor, str) or not valor.strip() for valor in registro[campo]):
            raise ErrorRegistro(f"{campo} contiene un valor vacío o que no es texto.")

    creditos = registro["creditos"]
    if not isinstance(creditos, dict):
        raise ErrorRegistro("Los créditos y horas deben declararse como campos.")
    campos_creditos = ("hc", "hl", "ht", "hpc", "hcl", "cr")
    faltan_creditos = [campo for campo in campos_creditos if campo not in creditos]
    if faltan_creditos:
        raise ErrorRegistro("Faltan créditos u horas: " + ", ".join(faltan_creditos))

    if not isinstance(registro["estudiantes"], list):
        raise ErrorRegistro("Los estudiantes deben declararse como lista.")
    if not isinstance(registro["responsables"], list):
        raise ErrorRegistro("Los responsables deben declararse como lista.")
    if len(registro["estudiantes"]) > 6:
        raise ErrorRegistro("La plantilla admite como máximo seis estudiantes por registro.")
    if len(registro["responsables"]) > 4:
        raise ErrorRegistro("La plantilla admite como máximo cuatro responsables por registro.")

    for i, estudiante in enumerate(registro["estudiantes"], start=1):
        if not isinstance(estudiante, dict):
            raise ErrorRegistro(f"Estudiante {i}: debe declararse como campos.")
        faltan = [campo for campo in ("matricula", "nombre", "programa") if not estudiante.get(campo)]
        if faltan:
            raise ErrorRegistro(f"Estudiante {i}: faltan " + ", ".join(faltan))
    for i, responsable in enumerate(registro["responsables"], start=1):
        if not isinstance(responsable, dict):
            raise ErrorRegistro(f"Responsable {i}: debe declararse como campos.")
        faltan = [campo for campo in ("nombre", "adscripcion", "ciudad") if not responsable.get(campo)]
        if faltan:
            raise ErrorRegistro(f"Responsable {i}: faltan " + ", ".join(faltan))

    for i, actividad in enumerate(registro["actividades"], start=1):
        if not isinstance(actividad, dict) or not actividad.get("descripcion"):
            raise ErrorRegistro(f"Actividad {i}: falta descripcion.")
        periodo = actividad.get("periodo_calendario")
        semanas = "inicio_semana" in actividad or "fin_semana" in actividad
        if periodo and semanas:
            raise ErrorRegistro(
                f"Actividad {i}: declara periodo_calendario y semanas al mismo tiempo."
            )
        if periodo:
            if not isinstance(calendario.periodos.get(periodo), Periodo):
                raise ErrorRegistro(
                    f"Actividad {i}: periodo inexistente en el calendario {registro['ciclo']}: "
                    f"{periodo}"
                )
            continue
        faltan = [campo for campo in ("inicio_semana", "fin_semana") if campo not in actividad]
        if faltan:
            raise ErrorRegistro(f"Actividad {i}: faltan " + ", ".join(faltan))
        inicio, fin = actividad["inicio_semana"], actividad["fin_semana"]
        if (
            isinstance(inicio, bool)
            or isinstance(fin, bool)
            or not isinstance(inicio, int)
            or not isinstance(fin, int)
            or not 1 <= inicio <= fin <= calendario.total_semanas
        ):
            raise ErrorRegistro(
                f"Actividad {i}: el rango de semanas debe cumplir "
                f"1 <= inicio <= fin <= {calendario.total_semanas}."
            )


def renderizar(entrada: Path, salida: Path) -> Path:
    datos = yaml.safe_load(Path(entrada).read_text(encoding="utf-8")) or {}
    _validar(datos)
    registro = datos["registro"]
    salida = Path(salida)
    salida.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix=f".{salida.stem}-", dir=salida.parent) as temporal:
        trabajo = Path(temporal) / salida.name
        copia_de_trabajo(PLANTILLA, trabajo)
        documento = Document(trabajo)

        generales = documento.tables[0]
        for fila, campo in enumerate(
            ("unidad_academica", "periodo_estudio", "tipo_modalidad", "clave", "nombre_modalidad")
        ):
            _campo(generales, fila, 1, registro[campo])

        creditos = documento.tables[1]
        for columna, campo in (
            (1, "hc"),
            (3, "hl"),
            (5, "ht"),
            (7, "hpc"),
            (9, "hcl"),
            (11, "cr"),
        ):
            _campo(creditos, 1, columna, registro["creditos"][campo])

        plan = documento.tables[2]
        _campo(plan, 0, 1, registro["plan_estudios"])
        _campo(plan, 0, 3, registro["etapa_formacion"])

        programas = documento.tables[3]
        _campo(programas, 1, 0, "\n".join(registro["programas"]))

        estudiantes = documento.tables[4]
        for fila, estudiante in enumerate(registro["estudiantes"], start=2):
            _campo(estudiantes, fila, 0, estudiante["matricula"])
            _campo(estudiantes, fila, 1, estudiante["nombre"])
            _campo(estudiantes, fila, 2, estudiante["programa"])
        _recortar_filas(estudiantes, 2 + len(registro["estudiantes"]))

        responsables = documento.tables[5]
        for fila, responsable in enumerate(registro["responsables"], start=2):
            _campo(responsables, fila, 0, responsable["nombre"])
            _campo(responsables, fila, 1, responsable["adscripcion"])
            _campo(responsables, fila, 2, responsable["ciudad"])
        _recortar_filas(responsables, 2 + len(registro["responsables"]))

        firmas = registro.get("firmas") or {}
        pie = documento.tables[6]
        for columna, campo in ((1, "coordinacion_extension"), (3, "direccion_subdireccion")):
            if firmas.get(campo):
                _campo(pie, 1, columna, firmas[campo])

        _rellenar_cuerpo(documento, registro)
        documento.save(trabajo)
        trabajo.replace(salida)
    return salida


def main(argv: list[str]) -> int:
    if len(argv) != 3:
        print("Uso: python src/render_registro_modalidades.py <registro.yaml> <salida.docx>", file=sys.stderr)
        return 2
    try:
        print(renderizar(Path(argv[1]), Path(argv[2])))
    except (ErrorRegistro, OSError, yaml.YAMLError) as error:
        print(f"Error: {error}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
