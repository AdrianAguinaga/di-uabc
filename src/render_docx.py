"""Rellena la plantilla CIAD real con los datos de `curso.yaml`.

No reconstruye el documento: **copia** la plantilla registrada y escribe sobre la copia,
clonando los elementos que ya trae para que el formato institucional —logo, estilos, bordes,
tipografía— sobreviva intacto.

Dos técnicas sostienen todo el módulo:

  1. **Prototipos.** En vez de crear párrafos y filas desde cero, se clona uno del propio
     documento con `deepcopy` y se le cambia el contenido. Así hereda todo el formato sin
     que el código tenga que conocerlo.
  2. **Runs, nunca `paragraph.text`.** Asignar `.text` colapsa los runs y destruye el
     formato. Se añaden runs heredando el `rPr` del que había (ver `estilo.py`).

Un documento por grupo:
    python src/render_docx.py cursos/2026-2/39056-big-data/curso.yaml
"""

from __future__ import annotations

import sys
from copy import deepcopy
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

import calendario
import estilo
import modelo
import plantillas
from modelo import Config, Curso, Grupo, Meta

RAIZ = Path(__file__).resolve().parent.parent

# Marcadores por los que se localiza cada pieza. Se busca por TEXTO: los índices se
# mueven entre versiones de la plantilla, el texto no.
M_SECCION_3 = "Sección 3."
M_META = "Meta 1.1."
M_META_2 = "Meta 1.2."
M_SEPARADOR = "------------------------"
M_VERSION = "Versión"
M_INSTRUCCIONES = "Para descargar instrucciones de llenado"

ORDINALES = [
    "Primero", "Segundo", "Tercero", "Cuarto", "Quinto", "Sexto",
    "Séptimo", "Octavo", "Noveno", "Décimo",
]


class ErrorRender(Exception):
    """La plantilla no tiene la forma que el renderizador espera."""


# -- utilidades de documento -------------------------------------------------


def _parrafos(doc) -> list[Paragraph]:
    return doc.paragraphs


def _buscar(doc, texto: str, desde: int = 0) -> int:
    """Índice del primer párrafo que contiene `texto`. Falla ruidosamente si no está."""
    for i, p in enumerate(doc.paragraphs[desde:], desde):
        if texto in p.text:
            return i
    raise ErrorRender(
        f"La plantilla no contiene el marcador {texto!r}. "
        f"¿Se actualizó sin revisar config/plantillas.yaml? "
        f"Corre `python src/plantillas.py verificar`."
    )


def _campo(doc, etiqueta: str) -> Paragraph:
    """El párrafo de la Sección 1 cuya etiqueta en negrita coincide."""
    for p in doc.paragraphs[:60]:
        if p.runs and p.runs[0].text.strip().rstrip(":") == etiqueta:
            return p
    raise ErrorRender(f"No hay campo «{etiqueta}» en la Sección 1 de la plantilla.")


def _anexar_valor(p: Paragraph, valor: str) -> None:
    """Añade el valor a un campo etiquetado, con exactamente un espacio de separación.

    La plantilla es inconsistente: unos campos terminan en `':'`, otros en `': '` y uno en
    `':  '`. Además, varios arrastran un run **vacío** al final, así que no basta con mirar
    el último run: hay que descartar primero los vacíos. Ese espacio de más se ve en el
    documento que acaba en manos de los alumnos.
    """
    if not (valor := str(valor).strip()):
        return
    while p.runs and not p.runs[-1].text:
        p.runs[-1]._r.getparent().remove(p.runs[-1]._r)
    if p.runs:
        p.runs[-1].text = p.runs[-1].text.rstrip()
    estilo.anexar(p, f" {valor}")


def _borrar(p: Paragraph) -> None:
    p._p.getparent().remove(p._p)


def _clonar_tras(prototipo: Paragraph, ancla) -> Paragraph:
    """Copia el prototipo y la inserta justo después de `ancla`. Devuelve la copia."""
    nuevo = deepcopy(prototipo._p)
    ancla.addnext(nuevo)
    return Paragraph(nuevo, prototipo._parent)


# -- Sección 1 ---------------------------------------------------------------


def _seccion_1(doc, curso: Curso, cfg: Config, grupo: Grupo) -> None:
    ident = curso.identificacion
    plantilla = cfg.plantilla(curso.modalidad)
    profesor = cfg.profesor(curso.profesor_id)

    # Portada. Los runs ya traen la etiqueta en negrita; solo falta el valor.
    _anexar_valor(doc.paragraphs[10], curso.nombre)
    _anexar_valor(doc.paragraphs[11], profesor["nombre"])

    # Título de la asignatura (párrafo suelto, sin etiqueta).
    estilo.reemplazar(doc.paragraphs[16], curso.nombre)

    campos = [
        ("Nivel académico", ident.get("nivel_academico", "")),
        ("Clave", f"{curso.clave}."),
        ("Programa Educativo", ident.get("programa_educativo", "")),
        ("Vigencia de la unidad de Aprendizaje", ident.get("plan_estudios", "")),
        ("Créditos y total de horas", ident.get("creditos_texto", "")),
        ("Etapa formativa", ident.get("etapa_formativa", "")),
        ("Carácter de la asignatura", ident.get("caracter", "")),
    ]
    for etiqueta, valor in campos:
        if valor:
            _anexar_valor(_campo(doc, etiqueta), valor)

    # La modalidad ya viene escrita en la plantilla; se comprueba, no se reescribe.
    p_mod = _campo(doc, "Modalidad Instruccional")
    esperada = plantilla["modalidad_instruccional"]
    if esperada.rstrip(".").lower() not in p_mod.text.lower():
        estilo.reemplazar(
            p_mod, [{"t": "Modalidad Instruccional", "enfasis": "etiqueta"},
                    {"t": f": {esperada}"}]
        )

    # Los cuatro bloques de contenido: el párrafo vacío que sigue a cada etiqueta.
    bloques = [
        ("Competencia general:", curso.contenido.get("competencia_general", "")),
        ("Propósito general:", curso.contenido.get("proposito_general", "")),
        ("Estrategia general de aprendizaje:", curso.contenido.get("estrategia_general", "")),
        ("Evidencias de desempeño:", curso.contenido.get("evidencias_desempeno", "")),
        ("Criterios de acreditación:", _texto_acreditacion(curso, cfg)),
        ("Criterios de evaluación:", curso.contenido.get("criterios_evaluacion", "")),
        ("Evaluación:", _texto_evaluacion(curso)),
    ]
    for etiqueta, texto in bloques:
        if texto:
            _escribir_tras(doc, etiqueta, texto)

    # Autores.
    i = _buscar(doc, "Programa de Unidad de Aprendizaje:")
    if autores := curso.contenido.get("autores_pua", ""):
        _anexar_valor(doc.paragraphs[i], autores)
    i = _buscar(doc, plantilla["etiqueta_autores"].split(":")[0])
    _anexar_valor(doc.paragraphs[i], profesor["nombre"] + ".")


def _escribir_tras(doc, etiqueta: str, texto: str) -> None:
    """Escribe el texto en los párrafos que siguen a una etiqueta de la Sección 1."""
    i = _buscar(doc, etiqueta)
    lineas = [l for l in str(texto).split("\n") if l.strip()]
    ancla = doc.paragraphs[i]._p
    hueco = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else None

    for n, linea in enumerate(lineas):
        if n == 0 and hueco is not None and not hueco.text.strip():
            estilo.reemplazar(hueco, linea)
            ancla = hueco._p
        else:
            p = _clonar_tras(hueco or doc.paragraphs[i], ancla)
            estilo.reemplazar(p, linea)
            ancla = p._p


def _texto_evaluacion(curso: Curso) -> str:
    lineas = [
        f"{r.etiqueta}: {_texto_rubro_evaluacion(r)}"
        + (f" ({r.detalle})" if r.detalle else "")
        for r in curso.rubros
    ]
    lineas.append(f"Total: {sum(r.porcentaje for r in curso.rubros):g}%")
    if curso.segundo_nivel is not None:
        lineas.extend([
            "La calificación final se integra de la siguiente manera:",
            f"{curso.segundo_nivel.promedio.etiqueta}: "
            f"{curso.segundo_nivel.promedio.porcentaje:g}%",
            f"{curso.segundo_nivel.ordinario.etiqueta}: "
            f"{curso.segundo_nivel.ordinario.porcentaje:g}%",
            f"Total: {curso.segundo_nivel.promedio.porcentaje + curso.segundo_nivel.ordinario.porcentaje:g}%",
        ])
    return "\n".join(lineas)


def _rubro(curso: Curso, id_: str):
    """El rubro de un aporte. Un curso que llegue aquí mal declarado no se disimula."""
    if (rubro := curso.rubro(id_)) is None:
        raise ErrorRender(f"La meta declara el rubro inexistente {id_!r}.")
    return rubro


def _valor_rubro(rubro, valor: float) -> str:
    """El valor en la unidad que declaró su rubro: puntos o porcentaje."""
    return f"{valor:g} pts" if rubro.unidad == "puntos" else f"{valor:g}%"


def _texto_rubro_evaluacion(rubro) -> str:
    """El porcentaje del rubro y, cuando corresponde, el total de puntos que reparte."""
    puntos = f" ({rubro.total:g} pts)" if rubro.unidad == "puntos" else ""
    return f"{rubro.porcentaje:g}%{puntos}"


def _valores_de_meta(curso: Curso, meta: Meta) -> list[tuple]:
    """Aportes de una meta en el orden del contrato, incluidos sus componentes.

    `Curso.aportes()` es la única definición de aporte para todos los consumidores. Llegar
    a los componentes desde aquí conservaría hoy el resultado, pero duplicaría una regla
    que R2 y R3 ya comparten.
    """
    return [
        (aporte, _rubro(curso, aporte.rubro))
        for aporte in curso.aportes()
        if aporte.meta is meta
    ]


def _texto_valor_meta(curso: Curso, meta: Meta) -> list[str]:
    """La declaración visible de la meta y de cada componente que le pertenece."""
    lineas = []
    for aporte, rubro in _valores_de_meta(curso, meta):
        valor = _valor_rubro(rubro, aporte.valor)
        if aporte.es_componente:
            lineas.append(f"{aporte.etiqueta} equivale a {valor} de {rubro.etiqueta}.")
        elif rubro.unidad == "puntos":
            lineas.append(
                f"La {meta.etiqueta.lower()} equivale a {valor} de {rubro.etiqueta}."
            )
        else:
            # Conserva al pie de la letra el texto de todos los DI a un solo nivel.
            lineas.append(
                f"La {meta.etiqueta.lower()} equivale al {valor} de tu calificación final."
            )
    return lineas


def _texto_acreditacion(curso: Curso, cfg: Config) -> str:
    """Los criterios de acreditación aplicables, con su fundamento legal."""
    lineas = []
    for c in cfg.politicas["acreditacion"]:
        if c.get("solo_si_practica") and not curso.practica:
            continue
        if (mods := c.get("modalidades")) and curso.modalidad not in mods:
            continue
        # Sin `profesores`, el criterio es de todos. Con la clave, solo de quien la
        # declara: así los criterios propios de un docente no se cuelan en el DI de otro.
        if (profes := c.get("profesores")) and curso.profesor_id not in profes:
            continue
        texto = c.get("texto") or c["plantilla"].format(exencion=curso.exencion_ordinario)
        if citas := c.get("citas"):
            texto += f" ({_citar(citas, cfg)})"
        lineas.append(" ".join(texto.split()))
    return "\n".join(lineas)


def _citar(claves: list[str], cfg: Config) -> str:
    """«Art. 70 del Estatuto Escolar UABC». Toda cita resuelve o revienta."""
    partes = []
    for clave in claves:
        art = cfg.articulo(clave)
        fuente = art.get("fuente", "Estatuto Escolar UABC")
        partes.append(f"Art. {art['articulo']} del {fuente.split(',')[0]}")
    return "; ".join(partes)


# -- Sección 2: la tabla -----------------------------------------------------


def _seccion_2(doc, curso: Curso, grupo: Grupo, cal: calendario.Calendario, cfg: Config) -> None:
    tabla = doc.tables[0]
    dividida = cfg.plantilla(curso.modalidad)["entrega_dividida"]

    # `tr_lst` reconsulta el XML en cada acceso: hay que congelarlo antes de borrar.
    filas = list(tabla._tbl.tr_lst)
    if len(filas) < 5:
        raise ErrorRender(f"La tabla de la plantilla solo tiene {len(filas)} filas.")

    encabezado, sep_unidad = filas[0], filas[1]
    proto_unidad = deepcopy(sep_unidad)
    proto_par = [deepcopy(f) for f in (filas[3], filas[4])] if dividida else [deepcopy(filas[3])]

    _celda(encabezado, 0, [
        [{"t": curso.nombre, "enfasis": "meta"}],
        [{"t": "Competencia: ", "enfasis": "etiqueta"},
         {"t": curso.contenido.get("competencia_general", "")}],
    ])

    # Se vacía todo lo que sigue al encabezado de columnas y se reconstruye.
    for tr in filas[3:]:
        tabla._tbl.remove(tr)
    tabla._tbl.remove(sep_unidad)

    ancla = filas[2]  # el encabezado de columnas
    for unidad in curso.unidades:
        fila_u = deepcopy(proto_unidad)
        _celda(fila_u, 0, f"Unidad {unidad.numero}: {unidad.nombre}"
               + (f"\n{unidad.competencia}" if unidad.competencia else ""))
        ancla.addnext(fila_u)
        ancla = fila_u

        for meta in curso.metas_de(unidad.numero):
            for tr in _filas_de_meta(curso, meta, proto_par, cal, dividida):
                ancla.addnext(tr)
                ancla = tr


def _evidencias(meta: Meta) -> list[str]:
    """Los nombres de evidencia que se imprimen para una meta: las suyas y las de sus
    componentes.

    Una meta sin componentes produce exactamente la lista de siempre — de eso depende que
    los cursos que no declaran el rasgo salgan carácter por carácter iguales.
    """
    return [e.nombre for e in meta.evidencias] + [
        c.evidencia.nombre for c in meta.componentes if c.evidencia
    ]


def _filas_de_meta(curso: Curso, meta: Meta, proto: list, cal, dividida: bool) -> list:
    """Las filas de la tabla que corresponden a una meta (dos si la entrega se divide)."""
    sesiones = meta.sesiones or []
    filas = []
    for n, plantilla_tr in enumerate(proto):
        tr = deepcopy(plantilla_tr)
        s = sesiones[n] if n < len(sesiones) else None
        ultima = n == len(proto) - 1

        _celda(tr, 0, [{"t": meta.etiqueta + ".", "enfasis": "meta"},
                       {"t": f" {meta.enunciado}"}] if n == 0 else "")
        _celda(tr, 1, str(min(meta.semanas)) if n == 0 else "")
        if dividida:
            if s and s.ambiente == "presencial":
                modo = (
                    "Virtual / Sincrónico"
                    if s.ambiente_efectivo == "virtual"
                    else "Presencial / Sincrónico"
                )
            else:
                modo = "Virtual / Asincrónico"
            _celda(tr, 2, modo)
        _celda(tr, 3 if dividida else 2,
               calendario.texto_fecha_corta(s.fecha) if s and s.fecha else "")
        _celda(tr, 4 if dividida else 3, s.actividad_tabla if s else "")
        _celda(tr, 5 if dividida else 4,
               ", ".join(_evidencias(meta)) if ultima else "")
        _celda(tr, 6 if dividida else 5,
               " / ".join(
                   _valor_rubro(rubro, aporte.valor)
                   for aporte, rubro in _valores_de_meta(curso, meta)
               ) if ultima else "")
        filas.append(tr)
    return filas


def _lineas(contenido) -> list:
    """Normaliza a «una entrada por párrafo». Cada entrada puede llevar tramos."""
    if isinstance(contenido, str):
        return contenido.split("\n")
    if isinstance(contenido, dict):
        return [contenido]
    if contenido and isinstance(contenido[0], (list, tuple)):
        return list(contenido)
    return [contenido]


def _celda(tr, indice: int, contenido) -> None:
    """Escribe en la celda `indice` de una fila, respetando su formato.

    Cada `<w:tc>` debe terminar con al menos un `<w:p>` o Word pide reparar el archivo,
    así que se reutiliza el párrafo que ya trae en vez de vaciarla entera.
    """
    celdas = tr.findall(qn("w:tc"))
    if indice >= len(celdas):
        return
    tc = celdas[indice]
    parrafos = tc.findall(qn("w:p"))
    if not parrafos:
        return

    for extra in parrafos[1:]:
        tc.remove(extra)
    molde = deepcopy(parrafos[0])
    p = Paragraph(parrafos[0], None)

    lineas = _lineas(contenido)
    estilo.reemplazar(p, lineas[0])
    for linea in lineas[1:]:
        nuevo = deepcopy(molde)
        p._p.addnext(nuevo)
        p = Paragraph(nuevo, None)
        estilo.reemplazar(p, linea)


def _sin_vmerge(tr) -> None:
    """Quita la fusión vertical de una fila reutilizada fuera de la Sección 2."""
    for tc in tr.findall(qn("w:tc")):
        tcpr = tc.get_or_add_tcPr()
        if (vmerge := tcpr.find(qn("w:vMerge"))) is not None:
            tcpr.remove(vmerge)


def _tabla_rubrica(doc, rubrica, ancla) -> object:
    """Inserta la rúbrica clonando la tabla y filas de la propia plantilla CIAD.

    La plantilla no trae una rúbrica. Se duplica su tabla de plan de actividades para
    heredar bordes, tipografía y anchuras registradas, y se reutiliza la fila presencial
    como modelo de cada renglón. No se construye ningún elemento OOXML desde cero.
    """
    origen = doc.tables[0]
    filas = list(origen._tbl.tr_lst)
    if len(filas) < 4:
        raise ErrorRender("La tabla de la plantilla no tiene filas suficientes para la rúbrica.")

    nueva = deepcopy(origen._tbl)
    ancla.addnext(nueva)
    tabla = Table(nueva, doc._body)
    for tr in list(tabla._tbl.tr_lst):
        tabla._tbl.remove(tr)

    def agregar(prototipo):
        fila = deepcopy(prototipo)
        _sin_vmerge(fila)
        tabla._tbl.append(fila)
        return fila

    titulo = agregar(filas[2])
    _celda(titulo, 0, [{"t": "Rúbrica de evaluación", "enfasis": "etiqueta"}])

    def datos(concepto: str, puntos: float | str, descripcion: str, enfasis="normal"):
        fila = agregar(filas[3])
        # La rejilla de la plantilla tiene siete columnas. Se conservan sus anchos: las dos
        # primeras dan espacio al concepto, la tercera al puntaje y las cuatro restantes a
        # la descripción, como una tabla CIAD de tres campos.
        renglon = tabla.rows[-1]
        renglon.cells[0].merge(renglon.cells[1])
        renglon.cells[3].merge(renglon.cells[-1])
        _celda(fila, 0, [{"t": concepto, "enfasis": enfasis}])
        _celda(fila, 1, [{"t": f"{puntos:g}" if isinstance(puntos, (int, float)) else puntos,
                          "enfasis": enfasis}])
        _celda(fila, 2, [{"t": descripcion, "enfasis": enfasis}])

    datos("Concepto", "Puntos", "Descripción", "etiqueta")
    for fila in rubrica.filas:
        datos(fila.concepto, fila.puntos, fila.descripcion)
    datos("Puntos totales", rubrica.total, "", "etiqueta")
    return tabla._tbl


# -- Sección 3: los bloques de meta ------------------------------------------


def _seccion_3(doc, curso: Curso, cfg: Config) -> dict:
    plantilla = cfg.plantilla(curso.modalidad)
    i_meta = _buscar(doc, M_META, _buscar(doc, M_SECCION_3))
    i_sep = _buscar(doc, M_SEPARADOR, i_meta)

    # Biblioteca de prototipos tomada del propio documento. Cada uno conserva su pPr.
    ps = doc.paragraphs
    proto = {
        "titulo": deepcopy(ps[i_meta]._p),
        "encabezado": deepcopy(ps[i_meta + 2]._p),   # ► ¿Qué voy a aprender?
        "normal": deepcopy(ps[i_meta + 6]._p),       # Carácter de la actividad: …
        "paso": deepcopy(ps[i_meta + 8]._p),         # Primero.
        "valor": deepcopy(ps[i_sep - 2]._p),         # Esta actividad tiene valor…
        "separador": deepcopy(ps[i_sep]._p),
    }
    encabezados = [ps[i].text.strip() for i in (i_meta + 2, i_meta + 5, i_meta + 19, i_meta + 22)]

    # Se borran los dos bloques de la plantilla y se reconstruyen uno por meta.
    i_fin = _buscar(doc, M_VERSION, i_sep)
    for p in doc.paragraphs[i_meta:i_fin]:
        _borrar(p)

    ancla = doc.paragraphs[_buscar(doc, M_VERSION)]._p
    ancla = ancla.getprevious() if ancla.getprevious() is not None else ancla
    for meta in curso.metas:
        ancla = _bloque_de_meta(doc, curso, meta, proto, encabezados, plantilla, ancla)
    return proto


def _bloque_de_meta(doc, curso: Curso, meta: Meta, proto: dict, encabezados, plantilla, ancla):
    """Arma el bloque de una meta clonando prototipos. Devuelve la última ancla."""

    def poner(clave: str, contenido, enfasis="normal"):
        nonlocal ancla
        nuevo = deepcopy(proto[clave])
        ancla.addnext(nuevo)
        ancla = nuevo
        estilo.reemplazar(Paragraph(nuevo, doc._body), contenido, enfasis)

    poner("titulo", [{"t": f"{meta.etiqueta}.", "enfasis": "meta"},
                     {"t": f" {meta.enunciado}"}])

    poner("encabezado", encabezados[0])
    for linea in meta.que_voy_a_aprender:
        poner("normal", linea)

    poner("encabezado", encabezados[1])
    poner("normal", [{"t": "Carácter de la actividad", "enfasis": "etiqueta"},
                     {"t": f": {meta.caracter}."}])

    n = 0
    for s in meta.sesiones:
        poner("normal", _rotulo_sesion(s, plantilla), "etiqueta")
        for paso in s.pasos:
            ordinal = ORDINALES[n] if n < len(ORDINALES) else f"Paso {n + 1}"
            poner("paso", [{"t": ordinal, "enfasis": "ordinal"},
                           {"t": ". "}, *_tramos_paso(paso)])
            n += 1

    poner("encabezado", encabezados[2])
    for s in meta.sesiones:
        if s.fecha and s.ambiente == "virtual":
            nombres = ", ".join(e.recurso or e.nombre for e in meta.evidencias)
            poner("normal", [
                {"t": nombres or meta.etiqueta, "enfasis": "recurso"},
                {"t": f", a más tardar el {_fecha_larga(s.fecha)}."},
            ])

    poner("encabezado", encabezados[3])
    for linea in meta.reflexion:
        poner("normal", linea)
    if meta.criterios_evaluacion:
        poner("normal", "Criterios de evaluación", "etiqueta")
        for c in meta.criterios_evaluacion:
            poner("normal", c)
    for linea in _texto_valor_meta(curso, meta):
        poner("valor", linea, "valor")
    poner("separador", "------------------------")
    return ancla


def _rotulo_sesion(s, plantilla) -> str:
    fecha = _fecha_larga(s.fecha) if s.fecha else "(fecha por resolver)"
    if s.ambiente == "presencial":
        if s.ambiente_efectivo == "virtual":
            return f"En línea / sesión síncrona ({fecha})"
        return f"En clase / sesión síncrona ({fecha})"
    if plantilla["entrega_dividida"]:
        return f"Fuera de clase / actividad asincrónica (antes del {fecha})"
    return f"Actividad en plataforma (antes del {fecha})"


def _fecha_larga(f) -> str:
    return f"{calendario.DIAS[f.weekday()]} {calendario.texto_fecha(f)}"


def _tramos_paso(paso) -> list[dict]:
    """Un paso es texto plano o una lista de tramos con énfasis semántico."""
    if isinstance(paso, str):
        return [{"t": paso}]
    if isinstance(paso, dict):
        return [paso]
    return [{"t": x} if isinstance(x, str) else x for x in paso]


# -- Bloques fusionados del cierre -------------------------------------------


def _anexos(doc, curso: Curso, grupo: Grupo, cal, cfg: Config, proto: dict) -> None:
    """Criterios del curso, reglas de convivencia, versión y firma del jefe de grupo.

    Nada de esto viene en la plantilla CIAD: es la fusión que reproduce el ejemplo 961.
    """
    # Fuera el pie de instrucciones del CIAD: no va en un documento entregable.
    i = 0
    while True:
        try:
            i = _buscar(doc, M_INSTRUCCIONES)
        except ErrorRender:
            break
        for p in doc.paragraphs[i:]:
            _borrar(p)
        break

    i_version = _buscar(doc, M_VERSION)
    p_version = doc.paragraphs[i_version]
    estilo.reemplazar(p_version, f"Versión {curso.ciclo}.")

    # Los prototipos vienen de la Sección 3: alineados a la izquierda, como el cuerpo
    # del documento. El párrafo «Versión» está centrado y no sirve de molde.
    proto_h, proto_n = proto["encabezado"], proto["normal"]

    previo = p_version._p.getprevious()
    ancla = previo if previo is not None else p_version._p

    def poner(contenido, encabezado=False, enfasis="normal"):
        nonlocal ancla
        nuevo = deepcopy(proto_h if encabezado and proto_h is not None else proto_n)
        ancla.addnext(nuevo)
        ancla = nuevo
        estilo.reemplazar(Paragraph(nuevo, doc._body), contenido,
                          "etiqueta" if encabezado else enfasis)

    poner("Criterios de evaluación del curso", encabezado=True)
    poner("La evaluación del curso se distribuye de la siguiente manera:")
    for r in curso.rubros:
        detalle = f" ({r.detalle})" if r.detalle else ""
        poner([{"t": f"{r.etiqueta}: ", "enfasis": "etiqueta"},
               {"t": f"{_texto_rubro_evaluacion(r)}{detalle}"}])
    poner([{"t": "Total: ", "enfasis": "etiqueta"},
           {"t": f"{sum(r.porcentaje for r in curso.rubros):g}%"}])

    if curso.segundo_nivel is not None:
        poner("La calificación final se integra de la siguiente manera:")
        for nivel in (curso.segundo_nivel.promedio, curso.segundo_nivel.ordinario):
            poner([{"t": f"{nivel.etiqueta}: ", "enfasis": "etiqueta"},
                   {"t": f"{nivel.porcentaje:g}%"}])
        poner([{"t": "Total: ", "enfasis": "etiqueta"}, {
            "t": f"{curso.segundo_nivel.promedio.porcentaje + curso.segundo_nivel.ordinario.porcentaje:g}%"
        }])

    if curso.rubrica is not None:
        ancla = _tabla_rubrica(doc, curso.rubrica, ancla)

    for linea in _texto_acreditacion(curso, cfg).split("\n"):
        poner(linea)

    for nombre, etiqueta in (
        ("examenes_ordinarios", "Periodo de exámenes ordinarios y de competencias"),
        ("examenes_extraordinarios", "Periodo de exámenes extraordinarios"),
    ):
        if isinstance(p := cal.periodos.get(nombre), calendario.Periodo):
            poner([{"t": f"{etiqueta}: ", "enfasis": "etiqueta"}, {"t": f"{p.texto()}."}])

    poner("Reglas de convivencia en el aula", encabezado=True)
    poner("Para garantizar un ambiente de aprendizaje respetuoso y productivo, se "
          "establecen las siguientes reglas de convivencia:")
    n = 0
    for regla in cfg.politicas["convivencia"]:
        if (mods := regla.get("modalidades")) and curso.modalidad not in mods:
            continue
        n += 1
        texto = regla.get("texto") or regla["plantilla"].format(
            tolerancia=curso.tolerancia_minutos
        )
        poner([{"t": f"{n}. {regla['titulo']}. ", "enfasis": "etiqueta"},
               {"t": " ".join(texto.split())}])
        poner([{"t": "Sanción: ", "enfasis": "etiqueta"},
               {"t": " ".join(regla["sancion"].split())}])

    poner([{"t": "Consecuencias por incumplimiento: ", "enfasis": "etiqueta"},
           {"t": " ".join(cfg.politicas["consecuencias"].split())}])

    poner("Fundamento normativo", encabezado=True)
    for clave in curso.citas:
        art = cfg.articulo(clave)
        fuente = art.get("fuente", "Estatuto Escolar de la UABC, reforma 20 de mayo de 2021")
        poner([{"t": f"Artículo {art['articulo']} — {fuente}. ", "enfasis": "etiqueta"},
               {"t": " ".join(art["resumen"].split())}])

    # La firma va al final, después de «Versión».
    ancla = p_version._p
    for linea in (
        f"Grupo {grupo.numero}",
        "Enterado del contenido de este Diseño Instruccional, "
        "de los criterios de evaluación y de las reglas de convivencia:",
        "",
        "Firma Jefe Grupo ______________________________________________________________",
    ):
        nuevo = deepcopy(proto_n)
        ancla.addnext(nuevo)
        ancla = nuevo
        estilo.reemplazar(Paragraph(nuevo, doc._body), linea)


def _tiene(doc, texto: str) -> bool:
    return any(texto in p.text for p in doc.paragraphs)


# -- orquestación ------------------------------------------------------------


def generar(curso: Curso, grupo: Grupo, destino: Path, cfg: Config | None = None) -> Path:
    """Genera el `.docx` de un grupo. Parte de una copia fresca de la plantilla."""
    cfg = cfg or Config()
    cal = calendario.cargar(curso.ciclo)
    modelo.resolver_fechas(curso, grupo, cal)

    plantillas.copia_de_trabajo(curso.modalidad, destino)
    doc = docx.Document(str(destino))

    _seccion_1(doc, curso, cfg, grupo)
    _seccion_2(doc, curso, grupo, cal, cfg)
    proto = _seccion_3(doc, curso, cfg)
    _anexos(doc, curso, grupo, cal, cfg, proto)

    doc.save(str(destino))
    return destino


def generar_todos(ruta_curso: Path, cfg: Config | None = None) -> list[Path]:
    """Un documento por grupo, en `<directorio del curso>/salida/`."""
    ruta_curso = Path(ruta_curso)
    curso = modelo.cargar(ruta_curso)
    salida = ruta_curso.parent / "salida"
    return [
        generar(curso, g, salida / curso.nombre_archivo(g.numero, "docx"), cfg)
        for g in curso.grupos
    ]


def main(argv: list[str]) -> int:
    if len(argv) < 2:
        print("Uso: python src/render_docx.py <ruta a curso.yaml>", file=sys.stderr)
        return 2
    try:
        for ruta in generar_todos(Path(argv[1])):
            print(f"→ {ruta}")
    except (ErrorRender, modelo.ErrorModelo, plantillas.ErrorPlantilla,
            calendario.ErrorCalendario) as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
