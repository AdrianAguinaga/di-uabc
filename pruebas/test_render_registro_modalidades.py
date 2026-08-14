"""Pruebas del Registro de Modalidades de Acreditación Diversas."""

from __future__ import annotations

import shutil
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

import yaml
from docx import Document
from docx.oxml.ns import qn

RAIZ = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(RAIZ / "src"))

import plantillas  # noqa: E402
import render_registro_modalidades as registro  # noqa: E402


def datos_validos() -> dict:
    """Caso sintético inspirado en los ejemplos, sin datos personales."""
    return {
        "registro": {
            "unidad_academica": "Facultad de Prueba",
            "periodo_estudio": "2026-2",
            "tipo_modalidad": "Unidad de Aprendizaje por Asesoría Académica",
            "clave": "00000",
            "nombre_modalidad": "Unidad de Aprendizaje de Prueba",
            "plan_estudios": "2021-2",
            "etapa_formacion": "Terminal",
            "programas": ["Programa Educativo de Prueba, Plan 2021-2"],
            "creditos": {"hc": 2, "hl": "", "ht": 2, "hpc": "", "hcl": "", "cr": 6},
            "estudiantes": [
                {"matricula": "0000000", "nombre": "ESTUDIANTE DE PRUEBA", "programa": "00000"}
            ],
            "responsables": [
                {
                    "nombre": "DOCENTE DE PRUEBA",
                    "adscripcion": "Facultad de Prueba",
                    "ciudad": "Tijuana",
                }
            ],
            "ciclo": "2026-2",
            "justificacion": (
                "La unidad pertenece a un plan anterior y se cursará mediante asesoría académica."
            ),
            "competencias": "Competencia general copiada literalmente del programa oficial.",
            "actividades": [
                {
                    "inicio_semana": 1,
                    "fin_semana": 3,
                    "descripcion": "Analiza los contenidos de la primera unidad",
                    "producto": "Informe de resultados",
                },
                {
                    "inicio_semana": 13,
                    "fin_semana": 13,
                    "descripcion": "Integra el producto final",
                },
            ],
            "evaluacion_acreditacion": (
                "Primer examen parcial 20 %.\nSegundo examen parcial 20 %.\n"
                "Actividades 30 %.\nProducto final 30 %.\nTotal 100 %."
            ),
            "referencias": [
                "Referencia básica 1.",
                "Referencia básica 2.",
                "Referencia básica 3.",
            ],
            "firmas": {
                "coordinacion_extension": "COORDINACIÓN DE PRUEBA",
                "direccion_subdireccion": "DIRECCIÓN DE PRUEBA",
            },
        }
    }


class RenderRegistroModalidades(unittest.TestCase):
    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp(prefix="di-registro-"))
        self.entrada = self.tmp / "registro.yaml"
        self.salida = self.tmp / "salida" / "registro.docx"

    def tearDown(self):
        shutil.rmtree(self.tmp, ignore_errors=True)

    def escribir(self, datos: dict) -> None:
        self.entrada.write_text(
            yaml.safe_dump(datos, allow_unicode=True, sort_keys=False), encoding="utf-8"
        )

    def test_renderiza_asesoria_con_fechas_oficiales_sin_tocar_la_plantilla(self):
        self.escribir(datos_validos())
        antes = plantillas.sha256(plantillas.cargar()[registro.PLANTILLA].ruta)

        self.assertEqual(self.salida, registro.renderizar(self.entrada, self.salida))

        documento = Document(self.salida)
        texto = "\n".join(
            nodo.text or "" for nodo in documento.element.body.iter(qn("w:t"))
        )
        self.assertIn("Unidad de Aprendizaje por Asesoría Académica", texto)
        self.assertIn("Del 10 de agosto de 2026 al 29 de agosto de 2026", texto)
        self.assertIn("Del 3 de noviembre de 2026 al 7 de noviembre de 2026", texto)
        self.assertIn("Entrega: Informe de resultados", texto)
        self.assertIn("Segundo examen parcial 20 %", texto)
        self.assertEqual(3, len(documento.tables[4].rows))
        self.assertEqual(3, len(documento.tables[5].rows))
        self.assertEqual(antes, plantillas.sha256(plantillas.cargar()[registro.PLANTILLA].ruta))

    def test_periodo_de_estudio_debe_coincidir_con_el_calendario(self):
        datos = datos_validos()
        datos["registro"]["periodo_estudio"] = "2026-1"
        self.escribir(datos)

        with self.assertRaisesRegex(registro.ErrorRegistro, "debe coincidir"):
            registro.renderizar(self.entrada, self.salida)
        self.assertFalse(self.salida.exists())

    def test_rechaza_un_intervalo_invertido_sin_crear_documento(self):
        datos = datos_validos()
        datos["registro"]["actividades"][0].update(inicio_semana=4, fin_semana=2)
        self.escribir(datos)

        with self.assertRaisesRegex(registro.ErrorRegistro, "1 <= inicio <= fin <= 16"):
            registro.renderizar(self.entrada, self.salida)
        self.assertFalse(self.salida.exists())

    def test_rechaza_semana_fuera_del_ciclo_sin_crear_documento(self):
        datos = datos_validos()
        datos["registro"]["actividades"][0]["fin_semana"] = 17
        self.escribir(datos)

        with self.assertRaisesRegex(registro.ErrorRegistro, "1 <= inicio <= fin <= 16"):
            registro.renderizar(self.entrada, self.salida)
        self.assertFalse(self.salida.exists())

    def test_un_fallo_de_render_no_reemplaza_una_salida_anterior(self):
        self.escribir(datos_validos())
        self.salida.parent.mkdir(parents=True)
        self.salida.write_bytes(b"salida anterior")

        with patch.object(
            registro, "_rellenar_cuerpo", side_effect=registro.ErrorRegistro("fallo simulado")
        ):
            with self.assertRaisesRegex(registro.ErrorRegistro, "fallo simulado"):
                registro.renderizar(self.entrada, self.salida)
        self.assertEqual(b"salida anterior", self.salida.read_bytes())


if __name__ == "__main__":
    unittest.main(verbosity=2)
