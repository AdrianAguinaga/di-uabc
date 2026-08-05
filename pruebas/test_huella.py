"""Pruebas del instrumento de la no contaminación (`src/huella.py`).

Lo que estas pruebas defienden es que la huella sea **determinista**: si el texto
extraído de un `.docx` variara entre corridas, el instrumento de REQ-48 daría falsos
positivos y nadie volvería a mirarlo. Deliberadamente no generan ningún curso real —abrir dos cursos completos con el
generador rompería el tiempo de la suite (D-18). Todo se prueba contra un `.docx`
sintético armado en memoria y contra un MANIFIESTO.yaml sintético escrito en un
directorio temporal.
"""

from __future__ import annotations

import shutil
import sys
import tempfile
import types
import unittest
from pathlib import Path

import docx
import yaml

RAIZ = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(RAIZ / "src"))

import huella  # noqa: E402


def documento_sintetico(destino: Path) -> Path:
    """Un .docx mínimo con lo único que complica la extracción: una celda fusionada."""
    doc = docx.Document()
    doc.add_paragraph("Antes de la tabla")
    tabla = doc.add_table(rows=2, cols=2)
    tabla.cell(0, 0).text = "Meta 0. Encuadre del curso."
    tabla.cell(0, 1).text = "Presencial"
    tabla.cell(1, 1).text = "Virtual"
    tabla.cell(0, 0).merge(tabla.cell(1, 0))   # vMerge: la fila 2 no tiene texto propio
    doc.add_paragraph("Después de la tabla")
    doc.save(str(destino))
    return destino


class EnRegistroTemporal(unittest.TestCase):
    """Nunca toca pruebas/huellas.yaml de verdad."""

    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp(prefix="di-huella-"))
        self._registro = huella.REGISTRO
        huella.REGISTRO = self.tmp / "huellas.yaml"

    def tearDown(self):
        huella.REGISTRO = self._registro
        shutil.rmtree(self.tmp, ignore_errors=True)


class ExtraccionDeTexto(EnRegistroTemporal):
    def test_el_texto_extraido_es_determinista(self):
        """Extraer dos veces el mismo .docx da el mismo sha_texto: sin esto, el
        instrumento daría falsos positivos en cada corrida."""
        ruta = documento_sintetico(self.tmp / "a.docx")
        h1 = huella.sha_texto(huella.extraer_texto(ruta))
        h2 = huella.sha_texto(huella.extraer_texto(ruta))
        self.assertEqual(h1, h2)

    def test_una_celda_fusionada_no_duplica_su_texto(self):
        """Defiende contra el efecto de row.cells[i].text con vMerge, documentado en
        09-RESEARCH.md: duplicaría el texto de la celda de origen en cada fila."""
        ruta = documento_sintetico(self.tmp / "b.docx")
        texto = huella.extraer_texto(ruta)
        self.assertEqual(1, texto.count("Meta 0. Encuadre del curso."))

    def test_el_texto_respeta_el_orden_del_documento(self):
        """Defiende que se recorra `body` y no las listas separadas
        paragraphs/tables, que no conservan el orden real."""
        ruta = documento_sintetico(self.tmp / "c.docx")
        texto = huella.extraer_texto(ruta)
        i_antes = texto.index("Antes de la tabla")
        i_meta = texto.index("Meta 0.")
        i_despues = texto.index("Después de la tabla")
        self.assertLess(i_antes, i_meta)
        self.assertLess(i_meta, i_despues)

    def test_el_texto_de_la_tabla_lleva_todas_sus_celdas(self):
        ruta = documento_sintetico(self.tmp / "d.docx")
        texto = huella.extraer_texto(ruta)
        self.assertIn("Presencial", texto)
        self.assertIn("Virtual", texto)


class RegistroDeHuellas(EnRegistroTemporal):
    def test_el_registro_va_y_vuelve(self):
        documentos = {
            "2026-2:39056:961": {
                "ciclo": "2026-2", "clave": "39056", "grupo": "961",
                "texto_docx": "aa" * 32, "informe": "bb" * 32, "manifiesto": "cc" * 32,
                "registrado": "2026-08-05",
            }
        }
        huella.guardar(documentos)
        self.assertEqual(documentos, huella.cargar())

    def test_el_registro_lleva_su_cabecera_de_no_editar(self):
        huella.guardar({})
        texto = huella.REGISTRO.read_text(encoding="utf-8")
        self.assertTrue(texto.startswith("#"))
        self.assertIn("no lo edites a mano", texto)

    def test_sin_registro_cargar_devuelve_vacio(self):
        self.assertEqual({}, huella.cargar())

    def test_verificar_sin_registro_es_error_huella(self):
        """Debe fallar antes de generar nada: con el registro vacío, verificar()
        lanza ErrorHuella en su primera línea, sin llegar a _generar_control."""
        with self.assertRaises(huella.ErrorHuella) as ctx:
            huella.verificar()
        self.assertIn("registrar", str(ctx.exception))


MANIFIESTO_DE_PRUEBA = {
    "generado": "2026-08-05T10:00:00",
    "commit": "abc1234",
    "curso": {"archivo": "cursos/x/curso.yaml", "sha256": "aa" * 32},
    "evaluacion": {"esquema_id": "estandar-2026",
                   "rubros": [{"id": "tareas", "porcentaje": 40}]},
    "archivos": [{"ruta": "salida/DI.docx", "grupo": "961",
                  "sha256": "bb" * 32, "bytes": 233051}],
}


class FormaDelManifiesto(EnRegistroTemporal):
    def _escribir(self, datos: dict) -> Path:
        ruta = self.tmp / "MANIFIESTO.yaml"
        ruta.write_text(
            yaml.safe_dump(datos, allow_unicode=True, sort_keys=False, default_flow_style=False),
            encoding="utf-8",
        )
        return ruta

    def test_la_marca_de_tiempo_y_el_commit_no_cuentan(self):
        """Defiende que la huella no dé un falso positivo en cada corrida."""
        import copy

        a = copy.deepcopy(MANIFIESTO_DE_PRUEBA)
        b = copy.deepcopy(MANIFIESTO_DE_PRUEBA)
        b["generado"] = "2026-08-06T11:30:00"
        b["commit"] = "def5678-sucio"
        h1 = huella.forma_del_manifiesto(self._escribir(a))
        h2 = huella.forma_del_manifiesto(self._escribir(b))
        self.assertEqual(h1, h2)

    def test_el_sha256_y_los_bytes_de_los_archivos_no_cuentan(self):
        """Defiende contra el zip del .docx, que cambia siempre."""
        import copy

        a = copy.deepcopy(MANIFIESTO_DE_PRUEBA)
        b = copy.deepcopy(MANIFIESTO_DE_PRUEBA)
        b["archivos"][0]["sha256"] = "ff" * 32
        b["archivos"][0]["bytes"] = 999999
        h1 = huella.forma_del_manifiesto(self._escribir(a))
        h2 = huella.forma_del_manifiesto(self._escribir(b))
        self.assertEqual(h1, h2)

    def test_una_clave_nueva_cambia_la_forma(self):
        """La prueba que da sentido a las otras dos: si esto no cambiara, la
        función no vigilaría nada."""
        import copy

        a = copy.deepcopy(MANIFIESTO_DE_PRUEBA)
        b = copy.deepcopy(MANIFIESTO_DE_PRUEBA)
        b["nivel_2"] = {"promedio": 60, "ordinario": 40}
        h1 = huella.forma_del_manifiesto(self._escribir(a))
        h2 = huella.forma_del_manifiesto(self._escribir(b))
        self.assertNotEqual(h1, h2)

    def test_el_sha256_del_curso_si_cuenta(self):
        """Un curso distinto no puede pasar por el mismo manifiesto."""
        import copy

        a = copy.deepcopy(MANIFIESTO_DE_PRUEBA)
        b = copy.deepcopy(MANIFIESTO_DE_PRUEBA)
        b["curso"]["sha256"] = "dd" * 32
        h1 = huella.forma_del_manifiesto(self._escribir(a))
        h2 = huella.forma_del_manifiesto(self._escribir(b))
        self.assertNotEqual(h1, h2)

    def test_un_archivo_de_mas_cambia_la_forma(self):
        """Se vigila la estructura, no los valores: un segundo archivo cambia la
        forma aunque su sha256/bytes se hayan quitado."""
        import copy

        a = copy.deepcopy(MANIFIESTO_DE_PRUEBA)
        b = copy.deepcopy(MANIFIESTO_DE_PRUEBA)
        b["archivos"].append({"ruta": "salida/DI2.docx", "grupo": "962",
                               "sha256": "ee" * 32, "bytes": 111111})
        h1 = huella.forma_del_manifiesto(self._escribir(a))
        h2 = huella.forma_del_manifiesto(self._escribir(b))
        self.assertNotEqual(h1, h2)


class GeneracionDeControl(EnRegistroTemporal):
    """Ejercita `_generar_control`, lo único que toca archivos reales del repo.

    `generar.paquete` se sustituye por un doble ligero que solo escribe un
    MANIFIESTO.yaml, como hace el real, sin invocar Word: generar dos cursos
    completos aquí rompería el tiempo de la suite (D-18)."""

    def setUp(self):
        super().setUp()
        self._raiz = huella.RAIZ
        self._control = huella.CONTROL
        self._paquete_original = huella.generar.paquete
        huella.RAIZ = self.tmp
        curso_dir = self.tmp / "curso-control"
        curso_dir.mkdir()
        (curso_dir / "curso.yaml").write_text("x: 1\n", encoding="utf-8")
        huella.CONTROL = (("curso-control/curso.yaml", ("961",)),)
        self.manifiesto = curso_dir / "MANIFIESTO.yaml"

        def falso(ruta, pdf, grupos):
            ruta.parent.joinpath("MANIFIESTO.yaml").write_text(
                yaml.safe_dump(MANIFIESTO_DE_PRUEBA, allow_unicode=True, sort_keys=False),
                encoding="utf-8",
            )
            return types.SimpleNamespace(
                curso=types.SimpleNamespace(ciclo="2026-2", clave="39056"),
                informe=types.SimpleNamespace(texto=lambda: "informe de prueba"),
                archivos=[],
                manifiesto=self.manifiesto,
            )

        huella.generar.paquete = falso

    def tearDown(self):
        huella.RAIZ = self._raiz
        huella.CONTROL = self._control
        huella.generar.paquete = self._paquete_original
        super().tearDown()

    def test_verificar_no_deja_el_manifiesto_si_no_existia_antes(self):
        """WR-01: si el MANIFIESTO.yaml no existía antes de la corrida,
        restaurar_manifiesto=True debe borrarlo, no dejarlo escrito -- de lo
        contrario verificar() deja de ser de solo lectura sobre el repo (D-28)."""
        self.assertFalse(self.manifiesto.exists())
        huella._generar_control(restaurar_manifiesto=True)
        self.assertFalse(self.manifiesto.exists())

    def test_verificar_restaura_el_manifiesto_que_ya_existia(self):
        """La otra mitad del mismo `if`: si sí había un MANIFIESTO.yaml antes,
        debe quedar exactamente como estaba, no con lo recién generado."""
        self.manifiesto.write_text("previo: true\n", encoding="utf-8")
        huella._generar_control(restaurar_manifiesto=True)
        self.assertEqual("previo: true\n", self.manifiesto.read_text(encoding="utf-8"))


if __name__ == "__main__":
    unittest.main()
