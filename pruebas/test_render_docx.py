"""Pruebas del renderizado, contra el curso real de Big Data 2026-2.

El criterio de aceptación de la fase es reproducir la estructura de `ejemplos/961 (1).pdf`.
Estas pruebas la comprueban pieza por pieza sobre el `.docx` generado, y además defienden dos
cosas que ninguna inspección visual detecta:

  - que renderizar **no toca la plantilla**, y
  - que dos grupos difieren **solo** donde deben diferir.
"""

from __future__ import annotations

import shutil
import sys
import tempfile
import unittest
from pathlib import Path

import docx
from docx.oxml.ns import qn

RAIZ = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(RAIZ / "src"))

import calendario  # noqa: E402
import modelo  # noqa: E402
import plantillas  # noqa: E402
import render_docx  # noqa: E402

CURSO = RAIZ / "cursos" / "2026-2" / "39056-big-data" / "curso.yaml"


def texto(doc) -> str:
    """Todo el texto del documento: párrafos y tablas."""
    partes = [p.text for p in doc.paragraphs]
    partes += [
        "".join(t.text or "" for t in tc.iter(qn("w:t")))
        for tabla in doc.tables
        for tr in tabla._tbl.tr_lst
        for tc in tr.findall(qn("w:tc"))
    ]
    return "\n".join(partes)


class RenderizadoDeBigData(unittest.TestCase):
    """Se renderiza una vez para toda la clase: abrir Word es caro."""

    @classmethod
    def setUpClass(cls):
        cls.tmp = Path(tempfile.mkdtemp(prefix="di-render-"))
        cls.curso = modelo.cargar(CURSO)
        cls.cal = calendario.cargar(cls.curso.ciclo)
        cls.rutas = {}
        for g in cls.curso.grupos:
            destino = cls.tmp / cls.curso.nombre_archivo(g.numero, "docx")
            render_docx.generar(cls.curso, g, destino)
            cls.rutas[g.numero] = destino
        cls.doc = docx.Document(str(cls.rutas["961"]))
        cls.texto = texto(cls.doc)

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.tmp, ignore_errors=True)

    # -- la promesa del sistema ---------------------------------------------

    def test_renderizar_no_toca_la_plantilla(self):
        self.assertEqual([], plantillas.verificar())

    def test_conserva_el_formato_de_la_plantilla(self):
        """Si se hubiera reconstruido el documento, no quedaría ni la portada ni la tabla."""
        self.assertIn("Universidad Autónoma de Baja California", self.texto)
        self.assertIn("Centro de Investigación para el Aprendizaje Digital", self.texto)
        self.assertEqual(1, len(self.doc.tables))

    # -- Sección 1 ----------------------------------------------------------

    def test_identificacion_literal_del_pua(self):
        for esperado in (
            "Clave: 39056.",
            "Programa Educativo: Licenciado en Inteligencia de Negocios.",
            "Vigencia de la unidad de Aprendizaje: 2021-2.",
            "Créditos y total de horas: 6 créditos (HC:1 HL:4 HE:1), 96 horas en el semestre.",
            "Etapa formativa: Disciplinaria.",
            "Carácter de la asignatura: Obligatoria.",
            "Modalidad Instruccional: Semipresencial.",
        ):
            self.assertIn(esperado, self.texto, f"falta el campo: {esperado}")

    def test_no_deja_dobles_espacios_en_los_campos(self):
        self.assertNotIn("Aprendizaje:  2021-2", self.texto)

    def test_competencia_general_copiada_sin_parafrasear(self):
        self.assertIn(
            "Gestionar datos masivos mediante la aplicación de herramientas y estrategias",
            self.texto,
        )

    def test_evaluacion_con_los_tres_rubros_y_el_total(self):
        for esperado in ("Exámenes: 20%", "Tareas y actividades de clase: 40%",
                         "Proyecto final: 40%", "Total: 100%"):
            self.assertIn(esperado, self.texto)

    def test_el_profesor_firma_el_diseno(self):
        self.assertIn("Adrian Rodriguez Aguiñaga", self.texto)

    # -- Sección 2: la tabla ------------------------------------------------

    def test_una_fila_separadora_por_unidad(self):
        for u in self.curso.unidades:
            self.assertIn(f"Unidad {u.numero}: {u.nombre}", self.texto)

    def test_dos_subfilas_por_meta_mas_encabezados_y_unidades(self):
        esperadas = 2 + len(self.curso.unidades) + 2 * len(self.curso.metas)
        self.assertEqual(esperadas, len(self.doc.tables[0].rows))

    def test_la_rejilla_sigue_teniendo_siete_columnas(self):
        filas = self.doc.tables[0]._tbl.tr_lst
        datos = [tr for tr in filas if len(tr.findall(qn("w:tc"))) == 7]
        self.assertEqual(2 * len(self.curso.metas), len(datos))

    def test_las_dieciseis_semanas_estan_en_la_tabla(self):
        celdas = [
            "".join(t.text or "" for t in tc.iter(qn("w:t")))
            for tr in self.doc.tables[0]._tbl.tr_lst
            for tc in tr.findall(qn("w:tc"))
        ]
        for semana in range(1, self.cal.total_semanas + 1):
            self.assertIn(str(semana), celdas, f"falta la semana {semana}")

    def test_ninguna_fecha_cae_en_dia_de_suspension(self):
        """16 sep, 2 nov y 16 nov de 2026. La regla invariable número 6."""
        prohibidas = {
            calendario.texto_fecha_corta(s.fecha) for s in self.cal.suspensiones
        } | {calendario.texto_fecha(s.fecha) for s in self.cal.suspensiones}
        for f in prohibidas:
            self.assertNotIn(f, self.texto, f"hay una entrega en día de suspensión: {f}")

    def test_ninguna_fecha_queda_fuera_del_periodo_de_clases(self):
        for m in self.curso.metas:
            for s in m.sesiones:
                self.assertTrue(
                    self.cal.en_periodo_de_clases(s.fecha),
                    f"{m.etiqueta} entrega el {s.fecha}, fuera del periodo de clases",
                )

    # -- Sección 3: los bloques de meta -------------------------------------

    def test_un_bloque_por_meta_en_orden(self):
        posiciones = [self.texto.find(f"{m.etiqueta}. {m.enunciado}") for m in self.curso.metas]
        self.assertNotIn(-1, posiciones, "alguna meta no tiene bloque en la Sección 3")
        self.assertEqual(sorted(posiciones), posiciones, "las metas no salen en orden")

    def test_cada_bloque_termina_con_su_separador(self):
        self.assertEqual(len(self.curso.metas), self.texto.count("------------------------"))

    def test_los_pasos_llevan_ordinales_corridos_entre_lo_sincrono_y_lo_asincrono(self):
        """La meta 0 tiene 2 pasos presenciales y 2 virtuales: Primero…Cuarto, sin reiniciar."""
        for ordinal in ("Primero.", "Segundo.", "Tercero.", "Cuarto."):
            self.assertIn(ordinal, self.texto)

    def test_distingue_la_sesion_presencial_de_la_asincronica(self):
        self.assertIn("En clase / sesión síncrona (", self.texto)
        self.assertIn("Fuera de clase / actividad asincrónica (antes del ", self.texto)

    def test_cada_meta_declara_su_valor_porcentual(self):
        for m in self.curso.metas:
            esperado = f"La {m.etiqueta.lower()} equivale al {m.valor:g}% de tu calificación final."
            self.assertIn(esperado, self.texto, f"falta el valor de {m.etiqueta}")

    def test_la_reflexion_va_en_pasado(self):
        for linea in self.curso.metas[1].reflexion:
            self.assertIn(linea, self.texto)

    # -- bloques fusionados --------------------------------------------------

    def test_criterios_de_evaluacion_del_curso(self):
        self.assertIn("Criterios de evaluación del curso", self.texto)
        self.assertIn("igual o mayor a 80", self.texto)

    def test_reglas_de_convivencia_cada_una_con_su_sancion(self):
        self.assertIn("Reglas de convivencia en el aula", self.texto)
        self.assertEqual(8, self.texto.count("Sanción:"))

    def test_cita_los_articulos_que_obligan(self):
        for art in ("Artículo 65", "Artículo 66", "Artículo 68", "Artículo 70", "Artículo 71"):
            self.assertIn(art, self.texto, f"falta {art}")

    def test_la_unidad_practica_cita_el_articulo_74(self):
        self.assertIn("Artículo 74", self.texto)
        self.assertIn("no se evalúa en extraordinario", self.texto)

    def test_los_periodos_de_examen_salen_del_calendario(self):
        self.assertIn("Del 30 de noviembre al 8 de diciembre de 2026", self.texto)
        self.assertIn("Del 14 de diciembre al 17 de diciembre de 2026", self.texto)

    def test_bloque_de_firma_del_jefe_de_grupo(self):
        self.assertIn("Grupo 961", self.texto)
        self.assertIn("Firma Jefe Grupo", self.texto)

    def test_quita_el_pie_de_instrucciones_del_ciad(self):
        """Es una nota para quien llena la plantilla; no va en el documento entregable."""
        self.assertNotIn("Para descargar instrucciones de llenado", self.texto)

    def test_la_version_es_la_del_ciclo_no_la_de_la_plantilla(self):
        self.assertIn("Versión 2026-2.", self.texto)
        self.assertNotIn("Versión 2025-1.", self.texto)

    # -- los dos grupos ------------------------------------------------------

    def test_cada_grupo_lleva_su_propio_bloque_de_firma(self):
        t962 = texto(docx.Document(str(self.rutas["962"])))
        self.assertIn("Grupo 962", t962)
        self.assertNotIn("Grupo 961", t962)

    def test_las_fechas_presenciales_divergen_porque_el_horario_difiere(self):
        """961 va los martes y 962 los jueves: todas las fechas de clase cambian."""
        t962 = texto(docx.Document(str(self.rutas["962"])))
        self.assertIn("martes 11 de agosto", self.texto)
        self.assertIn("jueves 13 de agosto", t962)
        self.assertNotIn("martes 11 de agosto", t962)

    def test_las_entregas_coinciden_porque_ambos_entregan_en_sabado(self):
        t962 = texto(docx.Document(str(self.rutas["962"])))
        self.assertIn("sábado 15 de agosto", self.texto)
        self.assertIn("sábado 15 de agosto", t962)

    def test_el_contenido_pedagogico_es_el_mismo_en_ambos(self):
        t962 = texto(docx.Document(str(self.rutas["962"])))
        for m in self.curso.metas:
            self.assertIn(f"{m.etiqueta}. {m.enunciado}", t962)


class Determinismo(unittest.TestCase):
    def test_regenerar_produce_el_mismo_texto(self):
        curso = modelo.cargar(CURSO)
        tmp = Path(tempfile.mkdtemp(prefix="di-det-"))
        try:
            salidas = [
                texto(docx.Document(str(
                    render_docx.generar(curso, curso.grupos[0], tmp / f"v{n}.docx")
                )))
                for n in (1, 2)
            ]
            self.assertEqual(salidas[0], salidas[1])
        finally:
            shutil.rmtree(tmp, ignore_errors=True)


class CriteriosPropiosDelDocente(unittest.TestCase):
    """Cada docente tiene criterios de acreditación propios; no deben mezclarse.

    Los de `zra` se tomaron de su DI de Contabilidad Financiera (`ejemplos/`). Los de
    `ara` son los que ya producían Big Data y Patrones, y no deben moverse.
    """

    def setUp(self):
        self.cfg = modelo.Config()
        self.curso = modelo.cargar(CURSO)  # de `ara`

    def _acreditacion(self, profesor_id: str) -> str:
        self.curso.profesor_id = profesor_id
        return render_docx._texto_acreditacion(self.curso, self.cfg)

    def test_los_criterios_de_zra_no_salen_en_el_di_de_ara(self):
        texto_ara = self._acreditacion("ara")
        self.assertNotIn("trabajo final completo es requisito", texto_ara)
        self.assertNotIn("código de ética", texto_ara)
        self.assertNotIn("sella", texto_ara)

    def test_los_criterios_de_zra_salen_en_su_di(self):
        texto_zra = self._acreditacion("zra")
        self.assertIn("trabajo final completo es requisito", texto_zra)
        self.assertIn("código de ética", texto_zra)
        self.assertIn("95 %", texto_zra)

    def test_los_criterios_comunes_los_conservan_ambos(self):
        """El filtro añade, no reemplaza: lo que no lleva `profesores` es de todos."""
        for profesor in ("ara", "zra"):
            with self.subTest(profesor=profesor):
                self.assertIn("escala de 0 a 100", self._acreditacion(profesor))

    def test_zra_firma_con_su_titulo_y_su_correo(self):
        zra = self.cfg.profesor("zra")
        self.assertEqual("Dra. Zurisaddai Rubio Arriaga", zra["nombre"])
        self.assertEqual("rubio.zurisaddai@uabc.edu.mx", zra["correo"])
        self.assertEqual(90, zra["exencion_predeterminada"])


class ErroresDeEntrada(unittest.TestCase):
    def test_un_ciclo_sin_calendario_no_se_renderiza(self):
        curso = modelo.cargar(CURSO)
        curso.ciclo = "2099-1"
        with self.assertRaises(calendario.ErrorCalendario):
            render_docx.generar(curso, curso.grupos[0], Path(tempfile.gettempdir()) / "x.docx")

    def test_una_modalidad_sin_plantilla_no_se_renderiza(self):
        curso = modelo.cargar(CURSO)
        curso.modalidad = "hibrida"
        with self.assertRaises((plantillas.ErrorPlantilla, modelo.ErrorModelo)):
            render_docx.generar(curso, curso.grupos[0], Path(tempfile.gettempdir()) / "x.docx")


class EvidenciaDeComponente(unittest.TestCase):
    """La evidencia de un componente llega al documento; su valor también en la Fase 13.

    La segunda mitad es la que defiende REQ-48: Big Data no declara componentes y su
    documento debe salir carácter por carácter igual que antes de esta fase.
    """

    @classmethod
    def setUpClass(cls):
        cls.tmp = Path(tempfile.mkdtemp(prefix="di-componente-"))
        cls.base = cls.render(modelo.cargar(CURSO), "base")

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.tmp, ignore_errors=True)

    @classmethod
    def render(cls, curso, nombre):
        destino = cls.tmp / f"{nombre}.docx"
        render_docx.generar(curso, curso.grupos[0], destino)
        return texto(docx.Document(str(destino)))

    def test_una_lista_de_componentes_vacia_no_cambia_ni_un_caracter(self):
        """Es la mitad de REQ-48 que se puede probar sin generar el paquete completo."""
        curso = modelo.cargar(CURSO)
        for m in curso.metas:
            m.componentes = []
        self.assertEqual(self.base, self.render(curso, "vacios"))

    def test_la_evidencia_del_componente_se_concatena_a_las_de_su_meta(self):
        curso = modelo.cargar(CURSO)
        meta = next(m for m in curso.metas if m.evidencias)
        meta.componentes = [
            modelo.Componente(
                rubro="examenes", valor=15, etiqueta="Examen I", tipo="examen_parcial",
                evidencia=modelo.Evidencia(nombre="Examen I resuelto", tipo="examen"),
            )
        ]
        salida = self.render(curso, "con-componente")
        self.assertNotIn("Examen I resuelto", self.base)
        self.assertIn(f"{meta.evidencias[-1].nombre}, Examen I resuelto", salida)

    def test_un_componente_sin_evidencia_imprime_su_valor_sin_inventar_evidencia(self):
        """La Fase 13 visibiliza el componente aunque no tenga evidencia propia."""
        curso = modelo.cargar(CURSO)
        curso.metas[1].componentes = [
            modelo.Componente(rubro="examenes", valor=15, etiqueta="Examen I",
                              tipo="examen_parcial")
        ]
        salida = self.render(curso, "sin-evidencia")
        self.assertIn("Examen I equivale a 15% de Exámenes.", salida)
        self.assertNotIn("Examen I resuelto", salida)

    def test_un_curso_que_no_declara_puntos_conserva_su_columna_porcentual(self):
        self.assertIn("10%", self.base)


class RenderizadoDeLaUnidadDeclarada(unittest.TestCase):
    """Los rasgos de evaluación variables llegan al documento sin usar 38985.

    Es un curso sintético construido sobre Big Data: la declaración literal de Contabilidad
    Financiera permanece reservada para la Fase 14.
    """

    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp(prefix="di-unidad-real-"))
        self.addCleanup(shutil.rmtree, self.tmp, True)
        self.curso = modelo.cargar(CURSO)

        tareas = self.curso.rubro("tareas")
        tareas.unidad = "puntos"
        tareas.total = 150

        self.meta = next(m for m in self.curso.metas if m.id == "1.1")
        self.meta.valor = 10
        self.meta.componentes = [
            modelo.Componente(
                rubro="examenes", valor=15, etiqueta="Examen I", tipo="examen_parcial"
            )
        ]
        self.curso.segundo_nivel = modelo.SegundoNivel(
            promedio=modelo.Nivel(
                porcentaje=60,
                etiqueta="Valor del promedio antes del Examen Ordinario",
            ),
            ordinario=modelo.Nivel(
                porcentaje=40,
                etiqueta="Valor del examen Ordinario",
            ),
        )
        self.curso.exencion_contra = "promedio"
        self.curso.rubrica = modelo.Rubrica(
            meta=self.meta.id,
            total=100,
            filas=[
                modelo.FilaRubrica(
                    concepto="Portada",
                    puntos=2,
                    descripcion="Datos de identificación completos.",
                ),
                modelo.FilaRubrica(
                    concepto="Conclusión",
                    puntos=98,
                    descripcion="Integra los resultados del trabajo.",
                ),
            ],
        )

        destino = self.tmp / "unidad-real.docx"
        render_docx.generar(self.curso, self.curso.grupos[0], destino)
        self.doc = docx.Document(str(destino))
        self.salida = texto(self.doc)

    def test_imprime_puntos_y_componentes_en_la_meta_que_los_declara(self):
        valores = [
            "".join(t.text or "" for t in tc.iter(qn("w:t")))
            for tr in self.doc.tables[0]._tbl.tr_lst
            for tc in tr.findall(qn("w:tc"))
        ]
        self.assertIn("10 pts / 15%", valores)
        self.assertIn("La meta 1.1 equivale a 10 pts de Tareas y actividades de clase.", self.salida)
        self.assertIn("Examen I equivale a 15% de Exámenes.", self.salida)

    def test_imprime_los_dos_niveles_con_las_etiquetas_declaradas(self):
        for esperado in (
            "Valor del promedio antes del Examen Ordinario: 60%",
            "Valor del examen Ordinario: 40%",
            "La calificación final se integra de la siguiente manera:",
        ):
            self.assertIn(esperado, self.salida)

    def test_agrega_una_tabla_de_rubrica_clonada_de_la_tabla_del_documento(self):
        self.assertEqual(2, len(self.doc.tables))
        tabla = self.doc.tables[1]
        celdas = lambda fila: [
            "".join(t.text or "" for t in tc.iter(qn("w:t")))
            for tc in fila._tr.findall(qn("w:tc"))
        ]
        self.assertEqual(5, len(tabla.rows))
        self.assertEqual("Rúbrica de evaluación", celdas(tabla.rows[0])[0])
        self.assertEqual(["Concepto", "Puntos", "Descripción"], celdas(tabla.rows[1]))
        self.assertEqual("Portada", celdas(tabla.rows[2])[0])
        self.assertEqual("2", celdas(tabla.rows[2])[1])
        self.assertEqual("Datos de identificación completos.", celdas(tabla.rows[2])[2])
        self.assertEqual("Puntos totales", celdas(tabla.rows[-1])[0])
        self.assertEqual("100", celdas(tabla.rows[-1])[1])
        self.assertEqual(self.doc.tables[0]._tbl.tblPr.xml, tabla._tbl.tblPr.xml)
        self.assertEqual(self.doc.tables[0]._tbl.tblGrid.xml, tabla._tbl.tblGrid.xml)


if __name__ == "__main__":
    unittest.main(verbosity=2)
